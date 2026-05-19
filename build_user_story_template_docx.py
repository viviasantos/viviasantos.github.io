from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = r"C:\Users\ssant\Desktop\GitHub\viviasantos.github.io\assets\user-story-template.docx"

BG = "F3F1ED"
SURFACE = "FBFAF7"
SURFACE_2 = "F2EFE9"
BORDER = "D9D5CE"
TEXT = "111111"
MUTED = "5F5A53"
ACCENT = "67D98F"
ACCENT_INK = "2F8B57"


def set_font(run, name="Calibri", size=11, color=TEXT, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def para(p, before=0, after=0, line=1.15, align=WD_ALIGN_PARAGRAPH.LEFT):
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def width(cell, inches):
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tcw = tc_pr.find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        tc_pr.append(tcw)
    tcw.set(qn("w:w"), str(int(inches * 1440)))
    tcw.set(qn("w:type"), "dxa")


def margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = tc_mar.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def border_table(table, color=BORDER):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def fixed_table(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    while len(grid.gridCol_lst) > 0:
        grid.remove(grid.gridCol_lst[0])
    for w in widths:
        el = OxmlElement("w:gridCol")
        el.set(qn("w:w"), str(int(w * 1440)))
        grid.append(el)
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            width(cell, w)
            margins(cell)


def clear(cell):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def text_cell(cell, text, *, size=10.5, color=TEXT, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    clear(cell)
    p = cell.paragraphs[0]
    para(p, 0, 0, 1.15, align)
    r = p.add_run(text)
    set_font(r, size=size, color=color, bold=bold, italic=italic)


def prompt_cell(cell, text):
    text_cell(cell, text, size=10.25, color=MUTED, italic=True)


def label_block(doc, title, subtitle=None):
    p = doc.add_paragraph()
    para(p, 0, 4, 1.0)
    r = p.add_run(title)
    set_font(r, name="Calibri", size=9.5, color=ACCENT_INK, bold=True)
    if subtitle:
        p2 = doc.add_paragraph()
        para(p2, 0, 10, 1.2)
        r2 = p2.add_run(subtitle)
        set_font(r2, size=10.5, color=MUTED)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.header_distance = Inches(0.45)
    sec.footer_distance = Inches(0.45)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = RGBColor.from_string(TEXT)

    footer = sec.footer.paragraphs[0]
    para(footer, 0, 0, 1.0, WD_ALIGN_PARAGRAPH.CENTER)
    r = footer.add_run("Viviane Santos | Business Analyst | User Story Template")
    set_font(r, size=9, color=MUTED)

    p = doc.add_paragraph()
    para(p, 0, 4, 1.0)
    r = p.add_run("BUSINESS ANALYST WORKSHEET")
    set_font(r, size=9.5, color=ACCENT_INK, bold=True)

    p = doc.add_paragraph()
    para(p, 0, 6, 0.95)
    r = p.add_run("User Story Template")
    set_font(r, name="Georgia", size=26, color=TEXT, bold=True)

    p = doc.add_paragraph()
    para(p, 0, 10, 1.25)
    r = p.add_run("Use this form to draft one story at a time. Keep the prompts short, fill the fields, and delete the placeholder text before sending it to the team.")
    set_font(r, size=10.5, color=MUTED)

    meta = doc.add_table(rows=3, cols=4)
    fixed_table(meta, [1.05, 2.2, 1.05, 2.2])
    border_table(meta)
    meta_items = [
        ("Story ID", "[US-001]"),
        ("Owner", "[name]"),
        ("Priority", "[High / Med / Low]"),
        ("Status", "[Draft / Review / Final]"),
        ("Area", "[project / product / process]"),
        ("Date", "[dd / mm / yyyy]"),
    ]
    for i, (label, value) in enumerate(meta_items):
        row = i // 2
        col = (i % 2) * 2
        text_cell(meta.rows[row].cells[col], label, size=9.25, color=ACCENT_INK, bold=True)
        shade(meta.rows[row].cells[col], SURFACE_2)
        prompt_cell(meta.rows[row].cells[col + 1], value)
        shade(meta.rows[row].cells[col + 1], SURFACE)

    label_block(doc, "Story statement", "Fill the three prompts below. This is the core of the story.")
    story = doc.add_table(rows=3, cols=2)
    fixed_table(story, [1.25, 5.25])
    border_table(story)
    for i, (label, value) in enumerate((
        ("As a", "[user or persona]"),
        ("I want", "[action or need]"),
        ("So that", "[business value]"),
    )):
        text_cell(story.rows[i].cells[0], label, size=9.5, color=ACCENT_INK, bold=True)
        shade(story.rows[i].cells[0], SURFACE_2)
        prompt_cell(story.rows[i].cells[1], value)
        shade(story.rows[i].cells[1], SURFACE)

    label_block(doc, "Context", "Use short notes only. Add more detail in comments if needed.")
    context = doc.add_table(rows=3, cols=2)
    fixed_table(context, [1.25, 5.25])
    border_table(context)
    for i, (label, value) in enumerate((
        ("Problem", "[what is happening today?]"),
        ("Scope", "[what is included or excluded?]"),
        ("Dependencies", "[systems, approvals, data, people]"),
    )):
        text_cell(context.rows[i].cells[0], label, size=9.5, color=ACCENT_INK, bold=True)
        shade(context.rows[i].cells[0], SURFACE_2)
        prompt_cell(context.rows[i].cells[1], value)
        shade(context.rows[i].cells[1], SURFACE)

    label_block(doc, "Acceptance criteria", "Write each line as something the team can test.")
    criteria = doc.add_table(rows=6, cols=3)
    fixed_table(criteria, [0.45, 4.65, 1.4])
    border_table(criteria)
    headers = ["#", "Criterion", "Validation / proof"]
    for i, h in enumerate(headers):
        text_cell(criteria.rows[0].cells[i], h, size=9.25, color=TEXT, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(criteria.rows[0].cells[i], SURFACE_2)
    for i in range(1, 6):
        text_cell(criteria.rows[i].cells[0], str(i), size=10, color=ACCENT_INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        prompt_cell(criteria.rows[i].cells[1], f"[criterion {i}]")
        prompt_cell(criteria.rows[i].cells[2], "[how will you validate it?]")
        shade(criteria.rows[i].cells[0], SURFACE_2)
        shade(criteria.rows[i].cells[1], SURFACE)
        shade(criteria.rows[i].cells[2], SURFACE)

    label_block(doc, "Notes and risks", "Keep anything extra here so the main story stays clean.")
    notes = doc.add_table(rows=2, cols=2)
    fixed_table(notes, [1.25, 5.25])
    border_table(notes)
    for i, (label, value) in enumerate((
        ("Open questions", "[questions to resolve]"),
        ("Risks / notes", "[blockers, assumptions, or reminders]"),
    )):
        text_cell(notes.rows[i].cells[0], label, size=9.5, color=ACCENT_INK, bold=True)
        shade(notes.rows[i].cells[0], SURFACE_2)
        prompt_cell(notes.rows[i].cells[1], value)
        shade(notes.rows[i].cells[1], SURFACE)

    label_block(doc, "Definition of done", "Tick the boxes when the story is ready to move forward.")
    done = doc.add_table(rows=4, cols=2)
    fixed_table(done, [0.55, 5.95])
    border_table(done)
    done_items = [
        "Reviewed with stakeholders",
        "Acceptance criteria approved",
        "Testing approach understood",
        "Dependencies confirmed",
    ]
    for i, item in enumerate(done_items):
        text_cell(done.rows[i].cells[0], "□", size=13, color=ACCENT_INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(done.rows[i].cells[0], SURFACE_2)
        text_cell(done.rows[i].cells[1], item, size=10.5, color=TEXT)
        shade(done.rows[i].cells[1], SURFACE)

    p = doc.add_paragraph()
    para(p, 0, 0, 1.25)
    r = p.add_run("Tip: ")
    set_font(r, size=10, color=ACCENT_INK, bold=True)
    r = p.add_run("delete the bracketed prompts before you share the template.")
    set_font(r, size=10, color=MUTED)

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build()
    print(OUT_PATH)
