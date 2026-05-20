from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\ssant\Desktop\GitHub\viviasantos.github.io")
OUT = ROOT / "downloads"

TEXT = "111111"
MUTED = "5F5A53"
ACCENT = "67D98F"
ACCENT_INK = "2F8B57"
BG = "F3F1ED"
SURFACE = "FBFAF7"
SURFACE_2 = "F2EFE9"
BORDER = "D9D5CE"
TITLE = "Georgia"
BODY = "Calibri"


def rgb(value):
    return RGBColor.from_string(value)


def set_font(run, *, name=BODY, size=11, color=TEXT, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def para(p, before=0, after=0, line=1.15, align=WD_ALIGN_PARAGRAPH.LEFT):
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    b = tbl_pr.first_child_found_in("w:tblBorders")
    if b is None:
        b = OxmlElement("w:tblBorders")
        tbl_pr.append(b)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = b.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            b.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = tc_mar.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")


def fixed_table(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    while len(grid.gridCol_lst) > 0:
        grid.remove(grid.gridCol_lst[0])
    for w in widths:
        g = OxmlElement("w:gridCol")
        g.set(qn("w:w"), str(int(w * 1440)))
        grid.append(g)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            cell_margins(cell)


def style_doc(doc, footer_text):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.header_distance = Inches(0.45)
    sec.footer_distance = Inches(0.45)
    styles = doc.styles
    styles["Normal"].font.name = BODY
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = rgb(TEXT)
    footer = sec.footer.paragraphs[0]
    para(footer, 0, 0, 1.0, WD_ALIGN_PARAGRAPH.CENTER)
    r = footer.add_run(footer_text)
    set_font(r, size=9, color=MUTED)
    return sec


def title_block(doc, kicker, title, subtitle):
    p = doc.add_paragraph()
    para(p, 0, 4, 1.0)
    r = p.add_run(kicker.upper())
    set_font(r, name=BODY, size=9.5, color=ACCENT_INK, bold=True)

    p = doc.add_paragraph()
    para(p, 0, 6, 0.96)
    r = p.add_run(title)
    set_font(r, name=TITLE, size=26, color=TEXT, bold=True)

    p = doc.add_paragraph()
    para(p, 0, 10, 1.25)
    r = p.add_run(subtitle)
    set_font(r, size=10.5, color=MUTED)


def section_heading(doc, text, after=4):
    p = doc.add_paragraph()
    para(p, 0, after, 1.0)
    r = p.add_run(text)
    set_font(r, name=TITLE, size=18, color=TEXT, bold=True)
    return p


def label_line(doc, label, value):
    p = doc.add_paragraph()
    para(p, 0, 2, 1.12)
    r = p.add_run(f"{label}: ")
    set_font(r, size=10.5, color=ACCENT_INK, bold=True)
    r = p.add_run(value)
    set_font(r, size=10.5, color=MUTED)


def small_callout(doc, title, body):
    t = doc.add_table(rows=1, cols=1)
    fixed_table(t, [6.5])
    cell = t.cell(0, 0)
    shading(cell, SURFACE_2)
    p = cell.paragraphs[0]
    para(p, 0, 2, 1.15)
    r = p.add_run(title)
    set_font(r, size=10.25, color=ACCENT_INK, bold=True)
    p = cell.add_paragraph()
    para(p, 0, 0, 1.25)
    r = p.add_run(body)
    set_font(r, size=10.5, color=TEXT)
    borders(t)
    return t


def section_table(doc, rows, widths):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    fixed_table(table, widths)
    borders(table)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            shading(cell, SURFACE_2 if r_idx == 0 else SURFACE)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            para(p, 0, 0, 1.12)
            r = p.add_run(value)
            set_font(r, size=9.8 if r_idx == 0 else 10.25, color=TEXT, bold=(r_idx == 0))
            if c_idx == 0:
                r.font.color.rgb = rgb(ACCENT_INK)
    return table


def make_cover_meta(doc, rows):
    section_table(doc, rows, [1.35, 2.55, 1.35, 1.25])


def add_blank_lines(doc, count=1, size=0):
    for _ in range(count):
        p = doc.add_paragraph()
        para(p, size, size, 1.0)



def build_scrum_guide():
    doc = Document()
    style_doc(doc, "Viviane Santos | Business Analyst | Scrum Guide")
    title_block(
        doc,
        "Didactic guide",
        "Scrum Implementation Guide",
        "This guide shows how to organize a simple project to apply Scrum in practice, with clarity, focus, and incremental delivery.",
    )
    make_cover_meta(doc, [
        ("Focus", "Scrum, Kanban, and prioritization"),
        ("Format", "Practical guide in plain language"),
        ("Delivery", "Step-by-step + templates"),
        ("Use", "Portfolio, study, and class use"),
    ])
    add_blank_lines(doc, 1)
    small_callout(doc, "How to run this project", "Start from a simple problem, structure the backlog, select the sprint based on value and capacity, and close with review and retrospective.")

    section_heading(doc, "1. Overview")
    p = doc.add_paragraph()
    para(p, 0, 0, 1.35)
    r = p.add_run("In this project, theory is transformed into a practical sequence: problem, backlog, prioritization, sprint, execution, review, and retrospective.")
    set_font(r, size=10.8, color=TEXT)
    p = doc.add_paragraph()
    para(p, 0, 0, 1.35)
    r = p.add_run("The goal is to show the reasoning clearly, in a way that is useful for study and portfolio presentation.")
    set_font(r, size=10.8, color=MUTED)

    section_heading(doc, "2. Step by step")
    steps = [
        ("Define the problem and goal", "The work starts with a simple scenario and a SMART goal. This keeps the proposal concrete and avoids too much detail too early."),
        ("Create the initial backlog", "Next, epics and user stories are created using INVEST logic. Each item should be small, negotiable, valuable, and testable."),
        ("Prioritize with MoSCoW", "During prioritization, the backlog is divided into Must, Should, Could, and Won't to keep the sprint realistic and value-driven."),
        ("Plan the sprint", "During planning, the goal, team capacity, selected items, and main risks are defined before the sprint starts."),
        ("Execute with Kanban", "During execution, visual columns, work-in-progress limits, and visible blockers are used for the whole team."),
        ("Review and Retro", "During review, the increment is presented, feedback is gathered, and the retrospective closes the loop with lessons for the next sprint."),
    ]
    for title, body in steps:
        p = doc.add_paragraph()
        para(p, 0, 5, 1.15)
        r = p.add_run(title)
        set_font(r, name=TITLE, size=13, color=ACCENT_INK, bold=True)
        p = doc.add_paragraph()
        para(p, 0, 4, 1.28)
        r = p.add_run(body)
        set_font(r, size=10.5, color=TEXT)

    section_heading(doc, "3. Frameworks and methodologies")
    table = doc.add_table(rows=4, cols=2)
    fixed_table(table, [2.1, 4.4])
    borders(table)
    fw = [
        ("Scrum + Kanban", "Scrum organizes work into sprints and events. Kanban helps visualize flow and keep focus on tasks that are actually in progress."),
        ("SMART + INVEST", "These techniques make planning clearer: a well-defined goal, better stories, and prioritization without overload."),
        ("MoSCoW + PDCA", "MoSCoW supports prioritization and PDCA helps continuous improvement in the next iteration."),
        ("Incremental delivery", "Small, testable deliveries make review easier and reduce risk."),
    ]
    for i, (a, b) in enumerate(fw):
        shading(table.rows[i].cells[0], SURFACE_2)
        shading(table.rows[i].cells[1], SURFACE)
        for j, val in enumerate((a, b)):
            p = table.rows[i].cells[j].paragraphs[0]
            para(p, 0, 0, 1.18)
            r = p.add_run(val)
            set_font(r, size=10 if j == 0 else 10.25, color=ACCENT_INK if j == 0 else TEXT, bold=(j == 0))
            table.rows[i].cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    section_heading(doc, "4. How to present")
    p = doc.add_paragraph()
    para(p, 0, 0, 1.35)
    r = p.add_run("The best way to explain this material is to start with the problem, show the sprint sequence, and close with the learning.")
    set_font(r, size=10.8, color=TEXT)
    p = doc.add_paragraph()
    para(p, 0, 0, 1.35)
    r = p.add_run("This makes it possible to show method, clarity, and practice without relying on long explanations.")
    set_font(r, size=10.8, color=MUTED)

    section_heading(doc, "5. Final checklist")
    checklist = doc.add_table(rows=4, cols=2)
    fixed_table(checklist, [0.55, 5.95])
    borders(checklist)
    items = [
        "The problem comes before the solution.",
        "The flow is shown in a few steps.",
        "The downloads serve as support for class or portfolio use.",
        "The closing highlights the result or an observed improvement.",
    ]
    for i, item in enumerate(items):
        shading(checklist.rows[i].cells[0], SURFACE_2)
        shading(checklist.rows[i].cells[1], SURFACE)
        p = checklist.rows[i].cells[0].paragraphs[0]
        para(p, 0, 0, 1.0, WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run("?")
        set_font(r, size=12.5, color=ACCENT_INK, bold=True)
        p = checklist.rows[i].cells[1].paragraphs[0]
        para(p, 0, 0, 1.12)
        r = p.add_run(item)
        set_font(r, size=10.5, color=TEXT)

    doc.save(OUT / "scrum_implementation_guide.docx")


def build_backlog_template():
    doc = Document()
    style_doc(doc, "Viviane Santos | Business Analyst | Backlog Template")
    title_block(
        doc,
        "Work template",
        "Product Backlog",
        "This template organizes epics, stories, priority, effort, and acceptance criteria clearly.",
    )
    make_cover_meta(doc, [
        ("Fields", "Epic, story, priority, effort"),
        ("Format", "Editable Word table"),
        ("Use", "Planning and refinement"),
        ("Output", "Backlog ready for review"),
    ])
    add_blank_lines(doc, 1)
    small_callout(doc, "How to fill it in", "Write one line per story, keep the description short, and make the acceptance criteria explicit.")
    section_heading(doc, "Backlog")
    table = doc.add_table(rows=6, cols=6)
    fixed_table(table, [0.65, 1.45, 2.55, 0.85, 0.65, 1.35])
    borders(table)
    headers = ["ID", "Epic", "User Story", "Priority", "Effort", "Acceptance Criteria"]
    for c, h in enumerate(headers):
        cell = table.rows[0].cells[c]
        shading(cell, SURFACE_2)
        p = cell.paragraphs[0]
        para(p, 0, 0, 1.05, WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(h)
        set_font(r, size=9.3, color=ACCENT_INK, bold=True)
    rows = [
        ("US-001", "Registration", "As a user, I want to sign up so I can access the system.", "Must", "5", "Form validates required fields and creates the account successfully."),
        ("US-002", "Login", "As a user, I want to sign in so I can see my tasks.", "Must", "3", "User is authenticated and invalid credentials show an error message."),
        ("US-003", "Dashboard", "As a user, I want to see a dashboard with demand status.", "Should", "8", "Dashboard shows counts by status and the last update."),
        ("US-004", "Notifications", "As a user, I want to receive alerts when something changes.", "Could", "5", "Notification is triggered when the task changes status."),
        ("US-005", "Report", "As a manager, I want to export a report to track indicators.", "Could", "8", "File is exported in CSV or PDF format."),
    ]
    for i, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            cell = table.rows[i].cells[c]
            shading(cell, SURFACE if i % 2 else "FFFFFF")
            p = cell.paragraphs[0]
            para(p, 0, 0, 1.10, WD_ALIGN_PARAGRAPH.LEFT if c in (1, 2, 5) else WD_ALIGN_PARAGRAPH.CENTER)
            r = p.add_run(value)
            set_font(r, size=9.2 if c == 2 or c == 5 else 9.4, color=TEXT, bold=(c == 0))
            if c in (3, 4):
                r.font.color.rgb = rgb(ACCENT_INK)
        table.rows[i].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    section_heading(doc, "Notes")
    p = doc.add_paragraph()
    para(p, 0, 0, 1.25)
    r = p.add_run("Each item should be small enough to fit in a sprint, and the acceptance criteria avoid ambiguity.")
    set_font(r, size=10.5, color=MUTED)
    doc.save(OUT / "product_backlog_template.docx")


def build_sprint_planning():
    doc = Document()
    style_doc(doc, "Viviane Santos | Business Analyst | Sprint Planning")
    title_block(
        doc,
        "Planning template",
        "Sprint Planning",
        "This template aligns the goal, capacity, selected items, and risks before the sprint begins.",
    )
    make_cover_meta(doc, [
        ("Goal", "Define the sprint goal"),
        ("Format", "Single page with blocks"),
        ("Use", "Planning and tracking"),
        ("Output", "Clear and lightweight plan"),
    ])
    add_blank_lines(doc, 1)
    section_heading(doc, "Sprint Goal")
    p = doc.add_paragraph()
    para(p, 0, 0, 1.2)
    r = p.add_run("Describe in one sentence what this sprint needs to deliver.")
    set_font(r, size=10.5, color=MUTED, italic=True)
    box = doc.add_table(rows=1, cols=1)
    fixed_table(box, [6.5])
    borders(box)
    shading(box.cell(0, 0), SURFACE)
    p = box.cell(0, 0).paragraphs[0]
    para(p, 0, 0, 1.2)
    r = p.add_run("[write the sprint goal here]")
    set_font(r, size=10.5, color=TEXT)

    section_heading(doc, "Capacity")
    cap = section_table(doc, [
        ("Total availability", "[hours]"),
        ("Absences", "[vacation / meetings / blockers]"),
        ("Estimated capacity", "[points / hours]"),
    ], [2.2, 4.3])
    for i in range(3):
        shading(cap.rows[i].cells[0], SURFACE_2)
        shading(cap.rows[i].cells[1], SURFACE)

    section_heading(doc, "Selected items")
    sel = doc.add_table(rows=4, cols=4)
    fixed_table(sel, [2.0, 2.2, 1.0, 1.3])
    borders(sel)
    for c, h in enumerate(["Item", "Value", "Effort", "Owner"]):
        shading(sel.rows[0].cells[c], SURFACE_2)
        p = sel.rows[0].cells[c].paragraphs[0]
        para(p, 0, 0, 1.0, WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(h)
        set_font(r, size=9.3, color=ACCENT_INK, bold=True)
    for i in range(1, 4):
        for c in range(4):
            shading(sel.rows[i].cells[c], SURFACE if i % 2 else "FFFFFF")
            p = sel.rows[i].cells[c].paragraphs[0]
            para(p, 0, 0, 1.05)
            r = p.add_run("")
            set_font(r, size=9.5, color=TEXT)

    section_heading(doc, "Risks and dependencies")
    risk = doc.add_table(rows=3, cols=2)
    fixed_table(risk, [2.0, 4.3])
    borders(risk)
    for i, (label, prompt) in enumerate((
        ("Risk 1", "[describe the risk]"),
        ("Risk 2", "[describe the risk]"),
        ("External dependency", "[system, approval, or external team]"),
    )):
        shading(risk.rows[i].cells[0], SURFACE_2)
        shading(risk.rows[i].cells[1], SURFACE)
        p = risk.rows[i].cells[0].paragraphs[0]
        para(p, 0, 0, 1.1)
        r = p.add_run(label)
        set_font(r, size=9.8, color=ACCENT_INK, bold=True)
        p = risk.rows[i].cells[1].paragraphs[0]
        para(p, 0, 0, 1.1)
        r = p.add_run(prompt)
        set_font(r, size=9.8, color=MUTED, italic=True)
    doc.save(OUT / "sprint_planning_template.docx")


def build_definition_of_done():
    doc = Document()
    style_doc(doc, "Viviane Santos | Business Analyst | Definition of Done")
    title_block(
        doc,
        "Quality checklist",
        "Definition of Done",
        "This checklist helps decide when a story can be considered ready to move forward.",
    )
    make_cover_meta(doc, [
        ("Goal", "Avoid ambiguity"),
        ("Format", "Visual checklist"),
        ("Use", "Review and validation"),
        ("Output", "Clear criteria"),
    ])
    add_blank_lines(doc, 1)
    section_heading(doc, "Checklist")
    table = doc.add_table(rows=6, cols=2)
    fixed_table(table, [0.6, 5.9])
    borders(table)
    items = [
        "The item has been implemented.",
        "The acceptance criteria have been met.",
        "Relevant tests have passed.",
        "The review has been completed.",
        "Documentation has been updated, if needed.",
        "The item is ready for demonstration.",
    ]
    for i, item in enumerate(items):
        shading(table.rows[i].cells[0], SURFACE_2)
        shading(table.rows[i].cells[1], SURFACE if i % 2 == 0 else "FFFFFF")
        p = table.rows[i].cells[0].paragraphs[0]
        para(p, 0, 0, 1.0, WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run("?")
        set_font(r, size=12.5, color=ACCENT_INK, bold=True)
        p = table.rows[i].cells[1].paragraphs[0]
        para(p, 0, 0, 1.12)
        r = p.add_run(item)
        set_font(r, size=10.3, color=TEXT)
    section_heading(doc, "Use")
    p = doc.add_paragraph()
    para(p, 0, 0, 1.25)
    r = p.add_run("This checklist is completed at the end of the sprint to keep quality visible and consistent.")
    set_font(r, size=10.5, color=MUTED)
    doc.save(OUT / "definition_of_done.docx")


def build_retrospective():
    doc = Document()
    style_doc(doc, "Viviane Santos | Business Analyst | Retrospective")
    title_block(
        doc,
        "Continuous learning",
        "Retrospective Template",
        "This template records what worked, what got in the way, and what will be adjusted in the next sprint.",
    )
    make_cover_meta(doc, [
        ("Goal", "Learn from the cycle"),
        ("Format", "Questions + actions"),
        ("Use", "Retro and improvement"),
        ("Output", "Actions with owners"),
    ])
    add_blank_lines(doc, 1)
    section_heading(doc, "Reflection questions")
    table = doc.add_table(rows=4, cols=2)
    fixed_table(table, [2.15, 4.35])
    borders(table)
    prompts = [
        ("What worked well?", "[describe the strengths]"),
        ("What got in the way?", "[describe the blockers]"),
        ("What should we repeat?", "[practices to keep]"),
        ("What will I change?", "[improvement action]"),
    ]
    for i, (label, prompt) in enumerate(prompts):
        shading(table.rows[i].cells[0], SURFACE_2)
        shading(table.rows[i].cells[1], SURFACE if i % 2 == 0 else "FFFFFF")
        p = table.rows[i].cells[0].paragraphs[0]
        para(p, 0, 0, 1.12)
        r = p.add_run(label)
        set_font(r, size=9.8, color=ACCENT_INK, bold=True)
        p = table.rows[i].cells[1].paragraphs[0]
        para(p, 0, 0, 1.12)
        r = p.add_run(prompt)
        set_font(r, size=9.8, color=MUTED, italic=True)
    section_heading(doc, "Actions")
    actions = doc.add_table(rows=3, cols=3)
    fixed_table(actions, [2.9, 1.4, 2.2])
    borders(actions)
    for c, h in enumerate(("Action", "Owner", "Due date")):
        shading(actions.rows[0].cells[c], SURFACE_2)
        p = actions.rows[0].cells[c].paragraphs[0]
        para(p, 0, 0, 1.0, WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(h)
        set_font(r, size=9.3, color=ACCENT_INK, bold=True)
    for i in range(1, 3):
        for c in range(3):
            shading(actions.rows[i].cells[c], SURFACE if i % 2 else "FFFFFF")
            p = actions.rows[i].cells[c].paragraphs[0]
            para(p, 0, 0, 1.05)
            r = p.add_run("")
            set_font(r, size=9.5, color=TEXT)
    doc.save(OUT / "retrospective_template.docx")


def build_kanban():
    doc = Document()
    style_doc(doc, "Viviane Santos | Business Analyst | Kanban Board")
    title_block(
        doc,
        "Visual flow",
        "Kanban Board Template",
        "This board visualizes tasks, limits work in progress, and keeps the flow easy to follow.",
    )
    make_cover_meta(doc, [
        ("Columns", "To Do, Doing, Review, Done"),
        ("Format", "Editable board"),
        ("Use", "Daily execution"),
        ("Output", "Less blocking"),
    ])
    add_blank_lines(doc, 1)
    section_heading(doc, "Board")
    board = doc.add_table(rows=2, cols=4)
    fixed_table(board, [1.62, 1.62, 1.62, 1.64])
    borders(board)
    cols = ["To Do", "Doing", "Review", "Done"]
    for i, label in enumerate(cols):
        shading(board.rows[0].cells[i], SURFACE_2)
        p = board.rows[0].cells[i].paragraphs[0]
        para(p, 0, 0, 1.0, WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(label)
        set_font(r, size=9.8, color=ACCENT_INK, bold=True)
        shading(board.rows[1].cells[i], SURFACE if i % 2 == 0 else "FFFFFF")
        p = board.rows[1].cells[i].paragraphs[0]
        para(p, 0, 0, 1.4)
        r = p.add_run("[cards here]")
        set_font(r, size=9.8, color=MUTED, italic=True)
    section_heading(doc, "Simple rules")
    rules = doc.add_table(rows=4, cols=2)
    fixed_table(rules, [2.6, 3.9])
    borders(rules)
    rule_rows = [
        ("WIP limit", "Keep only a few tasks at the same time."),
        ("Daily update", "Review the board every day."),
        ("Visible blockers", "Do not hide impediments."),
        ("Clear transition", "Only move what is ready to advance."),
    ]
    for i, (a, b) in enumerate(rule_rows):
        shading(rules.rows[i].cells[0], SURFACE_2)
        shading(rules.rows[i].cells[1], SURFACE if i % 2 == 0 else "FFFFFF")
        p = rules.rows[i].cells[0].paragraphs[0]
        para(p, 0, 0, 1.12)
        r = p.add_run(a)
        set_font(r, size=9.8, color=ACCENT_INK, bold=True)
        p = rules.rows[i].cells[1].paragraphs[0]
        para(p, 0, 0, 1.12)
        r = p.add_run(b)
        set_font(r, size=9.8, color=TEXT)
    doc.save(OUT / "kanban_board_template.docx")


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    build_scrum_guide()
    build_backlog_template()
    build_sprint_planning()
    build_definition_of_done()
    build_retrospective()
    build_kanban()


if __name__ == "__main__":
    main()
